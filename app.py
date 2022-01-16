# %%
import os
import docx
import pandas as pd
from bs4 import BeautifulSoup
from numpy import quantile
from docx.shared import RGBColor
from docx.shared import Pt



def info():
    try:
        feedback = ' '.join(qd.find("div", class_="results-feedback-wrapper").text.split())[9:]
        return feedback
    except:
        return ""

def checkArray(array, i):
    try:
        help = array[i]
        return help
    except:
        return ""


column_names = ["Otazka", "Odpoved", "Info", "A", "B", "C", "D", "E", "F", "G", "H", "I", "J", "K"]

df = pd.DataFrame(columns = column_names)
l_files = os.listdir("data")
for n_files in l_files:
    with open('data/' + n_files, 'r') as f:
        contents = f.read()
        soup = BeautifulSoup(contents, 'lxml')
        for qd in soup.find_all(lambda tag: tag.name == 'div' and tag.get('class') == ['qd']):
            questions = []
            for sarow in qd.find_all("div", class_="sarow"):
                # questions.append(' '.join(sarow.text.split())[3:]) without A.
                questions.append(' '.join(sarow.text.split()))
            df = df.append({
                "Otazka": ' '.join(qd.find("div", class_="qsholder").text.split()),
                "Odpoved": ' '.join(qd.find("p", class_="chosen").text.split()).split("You chose: ")[0].split("Correct answer: ")[1],
                "Info": info(),
                "A": checkArray(questions, 0),
                "B": checkArray(questions, 1),
                "C": checkArray(questions, 2),
                "D": checkArray(questions, 3),
                "E": checkArray(questions, 4),
                "F": checkArray(questions, 5),
                "G": checkArray(questions, 6),
                "H": checkArray(questions, 7),
                "I": checkArray(questions, 8),
                "J": checkArray(questions, 9),
                "K": checkArray(questions, 10)
                }, 
                ignore_index=True )


df = df.drop_duplicates(subset=['Otazka'])
df.to_excel('test.xlsx', engine='xlsxwriter')  
display(df)
print(len(df["Otazka"].unique()))
mydoc = docx.Document()
mydoc.add_paragraph("This is first paragraph of a MS Word file.")
mydoc.save("test.docx")
    
def setParagraf(text):
    p = document.add_paragraph(text)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p

def setInfoParagraf(text):
    p = document.add_paragraph(text)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(7)

def setGreenParagraf(text):
    run = setParagraf("").add_run(text)
    font = run.font
    font.color.rgb = RGBColor(0x32, 0xCD, 0x32)

def setHeaderParagraf(text):
    if text.find("[b]True or False:[/b] ") != -1:
        text = text.replace("[b]True or False:[/b] ", "")
        text = text + " [b](True or False)[/b]"
    text = text.replace("[/b]", "")
    question = text.split("[b]")
    run = setParagraf("").add_run(question[0])
    run.font.size = Pt(15)
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run2 = p.add_run(question[1])
    run2.italic = True

def setAnswerParagraf(text):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(text)
    run.italic = True

def checkAnswer(answer, trueAnswer):
    if answer != "":
        if trueAnswer.find(answer[:2]) == -1:
            setParagraf("     " + answer)
        else:
            setGreenParagraf("     " + answer)

document = docx.Document()
number = 0
for index, row in df.iterrows():
    number = number + 1
    setHeaderParagraf(str(number) + ") " + row['Otazka'])
    setAnswerParagraf("Odpovede: " + row['Odpoved'])
    checkAnswer(row["A"], row['Odpoved'])
    checkAnswer(row["B"], row['Odpoved'])
    checkAnswer(row["C"], row['Odpoved'])
    checkAnswer(row["D"], row['Odpoved'])
    checkAnswer(row["E"], row['Odpoved'])
    checkAnswer(row["F"], row['Odpoved'])
    checkAnswer(row["G"], row['Odpoved'])
    checkAnswer(row["H"], row['Odpoved'])
    checkAnswer(row["I"], row['Odpoved'])
    checkAnswer(row["J"], row['Odpoved'])
    checkAnswer(row["K"], row['Odpoved'])

    setInfoParagraf(row["Info"])
    document.add_paragraph("---------------------------------")

document.save('demo1.docx')

# %%
